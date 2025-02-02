# -*- coding: utf-8 -*-
import argparse
import requests
import openai
from docx import Document
from tqdm import tqdm
import time
import os
import pandas as pd
import logging 

# Remplacez par vos clés API
DEEPL_API_KEY = os.environ.get("DEEPL_API_KEY")
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")

# Configurez l'accès à OpenAI
openai.api_key = OPENAI_API_KEY

# Configuration du logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)


def create_glossary(api_key, name, source_lang, target_lang, glossary_path):
    api_url = "https://api.deepl.com/v2/glossaries"
    with open(glossary_path, "r") as glossary_file:
        glossary_content = glossary_file.read()
    
    headers = {
        "Authorization": f"DeepL-Auth-Key {api_key}",
        "Content-Type": "application/x-www-form-urlencoded"
    }
    
    data = {
        "name": name,
        "source_lang": source_lang,
        "target_lang": target_lang,
        "entries_format": "csv",
        "entries": glossary_content
    }
    
    response = requests.post(api_url, headers=headers, data=data)
    response_data = response.json()
    
    if response.status_code in (200, 201) and "glossary_id" in response_data:
        glossary_id = response_data["glossary_id"]
        print("Glossary created successfully.")
        return glossary_id
    else:
        print(f"Failed to create glossary: {response.text}")
        raise Exception(f"Failed to create glossary: {response.text}")


# Section de traduction avec DeepL
def translate_docx_with_deepl(api_key, input_file_path, output_file_path, target_language, source_language, glossary_id=None):
    api_url = "https://api.deepl.com/v2/document"
    
    # En-têtes avec Bearer token
    headers = {
        "Authorization": f"Bearer {api_key}"
    }

    # Données pour la requête
    data = {
        "auth_key": api_key,  # Inclure explicitement l'auth_key
        "target_lang": target_language,
        "source_lang": source_language
    }
    if glossary_id:
        data["glossary_id"] = glossary_id

    # Journalisation pour débogage
    print(f"Uploading document with headers: {headers}")
    print(f"Data being sent (excluding file): {data}")

    # Téléversement du document
    with open(input_file_path, 'rb') as file:
        upload_response = requests.post(api_url, headers=headers, data=data, files={"file": file})

    print(f"Upload response status code: {upload_response.status_code}")
    print(f"Upload response content: {upload_response.text}")

    if upload_response.status_code != 200:
        raise Exception(f"Failed to upload document: {upload_response.text}")

    # Analyse de la réponse
    upload_data = upload_response.json()
    document_id = upload_data["document_id"]
    document_key = upload_data["document_key"]

    print("Document uploaded successfully.")
    print(f"Document ID: {document_id}, Document Key: {document_key}")

    # Suivi du statut
    status_url = f"{api_url}/{document_id}"
    while True:
        status_response = requests.post(status_url, headers=headers, data={"auth_key": api_key, "document_key": document_key})
        print(f"Status response: {status_response.status_code}, {status_response.text}")

        if status_response.status_code != 200:
            raise Exception(f"Failed to check translation status: {status_response.text}")

        status_data = status_response.json()
        if status_data["status"] == "done":
            print("Translation completed successfully.")
            break
        elif status_data["status"] == "error":
            raise Exception(f"Translation error: {status_data}")

    # Téléchargement du document traduit
    download_url = f"{api_url}/{document_id}/result"
    download_response = requests.post(download_url, headers=headers, data={"auth_key": api_key, "document_key": document_key})

    print(f"Download response status code: {download_response.status_code}")
    if download_response.status_code == 200:
        with open(output_file_path, "wb") as output_file:
            output_file.write(download_response.content)
        print(f"Translated document saved to: {output_file_path}")
    else:
        raise Exception(f"Failed to download translated document: {download_response.text}")




def convert_excel_to_csv(excel_path, csv_path):
    """
    Convertit un fichier Excel (.xlsx) en CSV pour DeepL.
    """
    df = pd.read_excel(excel_path, header=None)
    df.to_csv(csv_path, index=False, header=False)
    return csv_path


def read_glossary(glossary_path):
    glossary = {}
    doc = Document(glossary_path)
    for line in doc.paragraphs:
        if ":" in line.text:
            source, target = line.text.split(":", 1)
            glossary[source.strip()] = target.strip()
    return glossary


def process_paragraphs(paragraphs, glossary, language_level, source_language, target_language, model):
    """
    Envoie les paragraphes à ChatGPT pour amélioration de la traduction.
    """
    print(f"Sending the following paragraphs to ChatGPT:\n{paragraphs}\n")

    prompt = (
        f"Translate the following text from {source_language} to {target_language} "
        f"and improve its quality to match the '{language_level}' language level.\n"
        f"Use the glossary strictly when applicable: {glossary}.\n"
        f"Return only the improved translation, without any additional comments, explanations, or introductory phrases.\n\n"
    )

    for para in paragraphs:
        prompt += f"{para}\n\n"

    try:
        response = openai.ChatCompletion.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are a skilled translator and editor."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=2048,
            temperature=0.7,
        )
        return response["choices"][0]["message"]["content"].strip()
    except openai.error.RateLimitError as e:
        print(f"Rate limit reached: {e}. Adding delay before retrying.")
        time.sleep(15)
        return None
    except Exception as e:
        print(f"An error occurred with OpenAI API: {e}")
        return None


def improve_translation(input_file, glossary_path, output_file, language_level, source_language, target_language, group_size, model):
    doc = Document(input_file)
    glossary = read_glossary(glossary_path)
    output_doc = Document()
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    print(f"Loaded {len(paragraphs)} paragraphs from the translated document.")
    with tqdm(total=len(paragraphs), desc="Processing paragraphs") as pbar:
        for i in range(0, len(paragraphs), group_size):
            group = paragraphs[i : i + group_size]
            improved_text = process_paragraphs(group, glossary, language_level, source_language, target_language, model)
            if improved_text:
                output_doc.add_paragraph(improved_text)
            else:
                print(f"Skipping group {i // group_size + 1} due to an error.")
            pbar.update(len(group))
    output_doc.save(output_file)
    print(f"Improved document saved to: {output_file}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Translate and improve documents using DeepL and ChatGPT.")
    parser.add_argument("input_file", help="Path to the input .docx file.")
    parser.add_argument("translated_file", help="Path to save the translated .docx file.")
    parser.add_argument("improved_file", help="Path to save the improved .docx file.")
    parser.add_argument("source_language", help="Source language code (e.g., 'EN', 'FR').")
    parser.add_argument("target_language", help="Target language code (e.g., 'EN', 'FR').")
    parser.add_argument("language_level", help="Language level for improved translation (e.g., 'soutenu').")
    parser.add_argument("group_size", type=int, help="Number of paragraphs to process together.")
    parser.add_argument("--glossary_csv", help="Path to glossary CSV for DeepL.", default=None)
    parser.add_argument("--glossary_gpt", help="Path to glossary Word for ChatGPT.", default=None)
    parser.add_argument("--gpt_model", choices=["gpt-3.5-turbo", "gpt-4"], default="gpt-3.5-turbo", help="Choose the GPT model to use.")
    args = parser.parse_args()

    try:
        glossary_id = None
        if args.glossary_csv:
            glossary_id = create_glossary(
                api_key=DEEPL_API_KEY,
                name="MyGlossary",
                source_lang=args.source_language,
                target_lang=args.target_language,
                glossary_path=args.glossary_csv,
            )
        translate_docx_with_deepl(
            api_key=DEEPL_API_KEY,
            input_file_path=args.input_file,
            output_file_path=args.translated_file,
            target_language=args.target_language,
            source_language=args.source_language,
            glossary_id=glossary_id,
        )
        improve_translation(
            input_file=args.translated_file,
            glossary_path=args.glossary_gpt,
            output_file=args.improved_file,
            language_level=args.language_level,
            source_language=args.source_language,
            target_language=args.target_language,
            group_size=args.group_size,
            model=args.gpt_model,
        )
    except Exception as e:
        print(f"An error occurred: {e}")
